"""Tests for lib/documents.py (Task 17: reference document corpus).

Two tiers:
  * unit tests build a small synthetic .docx under tmp_path so the two
    style gaps (Properties: under Normal, .1 properties under CRM Dot One
    Property), the heading-path stack, table folding/skipping, and the
    chunk_document record shape can be exercised precisely and fast;
  * integration tests read the real cidoc_crm_version_7.3.2.docx (never
    write to it -- it is open in LibreOffice) and assert the invariants the
    task calls out: 244 declarations, E55/P2/E1 singletons, a real .1
    property surviving somewhere, the Minimality narrative surviving, and
    total narrative chars close to the measured 191,829.
"""

import docx
import pytest
from docx.enum.style import WD_STYLE_TYPE

from lib.config import PROJECT_ROOT, load_config
from lib.documents import build_documents, chunk_document, load_document, parse_docx

REAL_DOCX = PROJECT_ROOT / "sources" / "cidoc_crm_version_7.3.2.docx"

FAKE_ONTO = {
    "classes": {"E1": {}, "E55": {}, "E28": {}},
    "properties": {"P2": {}, "P3": {}, "P150": {}},
}


def _style(document, name):
    try:
        document.styles[name]
    except KeyError:
        document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    return name


def _p(document, text, style):
    document.add_paragraph(text, style=_style(document, style))


def make_synthetic_docx(tmp_path):
    """A small .docx exercising both style gaps, headings, and tables."""
    d = docx.Document()

    # Front matter: no heading is open yet, so this must be dropped.
    _p(d, "Cover page title", "Normal")

    # Nested heading path stack.
    d.add_heading("Modelling principles", level=1)
    d.add_heading("Minimality", level=2)
    _p(d, "The model itself is constructed as economically as possible.", "Normal")

    # Two headings back to back -> the first must be dropped as empty.
    d.add_heading("Empty Section", level=2)
    d.add_heading("Monotonicity", level=2)
    _p(d, "Monotonicity body text.", "Body Text")

    # A glossary-style table under a plain heading -- folds into narrative.
    d.add_heading("Terminology", level=2)
    _p(d, "Intro to terms.", "Normal")
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "class"
    table.cell(0, 1).text = "A category of items."
    table.cell(1, 0).text = "property"
    table.cell(1, 1).text = "A relation between two classes."

    # The hierarchy tables under this exact heading are redundant with the
    # XML and must be skipped even though a block is open.
    d.add_heading("Class & Property Hierarchies", level=1)
    d.add_heading("CIDOC CRM Class Hierarchy", level=2)
    _p(d, "See table below.", "Body Text")
    hier = d.add_table(rows=1, cols=2)
    hier.cell(0, 0).text = "E1"
    hier.cell(0, 1).text = "CRM Entity"

    # A long section that must split under chunking.
    d.add_heading("Big Section", level=1)
    _p(d, "word " * 1000, "Normal")

    # Declarations section.
    d.add_heading("CIDOC CRM Class Declarations", level=1)
    _p(d, "Classes are declared below.", "Body Text")

    _p(d, "E1 CRM Entity", "CRM Class Label")
    _p(d, "Scope note:", "CRM Description Label")
    _p(d, "E1 is the top class.", "CRM Scope Note Text")
    # Gap 1: "Properties:" under Normal, not CRM Description Label.
    _p(d, "Properties:", "Normal")
    _p(d, "P2 has type (is type of): E55 Type", "CRM Property of Entity")
    # Gap 2: the .1 property under its own dedicated style.
    _p(d, "(P3.1 has type: E55 Type)", "CRM Dot One Property")
    _p(d, "In first-order logic:", "CRM Description Label")
    _p(d, "E1(x)", "CRM First Order Logic")

    _p(d, "E55 Type", "CRM Class Label")
    _p(d, "Scope note:", "CRM Description Label")
    _p(d, "E55 comprises controlled-vocabulary terms.", "CRM Scope Note Text")
    _p(d, "In first-order logic:", "CRM Description Label")
    _p(d, "E55(x) => E28(x)", "CRM First Order Logic")

    _p(d, "P2 has type (is type of)", "CRM Property Label")
    _p(d, "Domain:", "CRM Description Label")
    _p(d, "E1 CRM Entity", "CRM Domain Range")

    path = tmp_path / "synthetic.docx"
    d.save(str(path))
    return path


@pytest.fixture(scope="module")
def synthetic_blocks(tmp_path_factory):
    tmp_path = tmp_path_factory.mktemp("docs")
    path = make_synthetic_docx(tmp_path)
    return parse_docx(path)


# --- unit tests on the synthetic fixture -----------------------------------


def test_front_matter_before_first_heading_is_dropped(synthetic_blocks):
    assert not any("Cover page title" in b["text"] for b in synthetic_blocks)


def test_heading_path_stack_tracks_nesting(synthetic_blocks):
    minimality = next(b for b in synthetic_blocks if b["heading"] == "Minimality")
    assert minimality["section_path"] == ["Modelling principles", "Minimality"]
    assert "economically" in minimality["text"]


def test_empty_section_between_two_headings_is_not_emitted(synthetic_blocks):
    assert not any(b["heading"] == "Empty Section" for b in synthetic_blocks)


def test_table_folds_into_the_open_narrative_block(synthetic_blocks):
    terminology = next(b for b in synthetic_blocks if b["heading"] == "Terminology")
    assert "A category of items." in terminology["text"]
    assert "A relation between two classes." in terminology["text"]


def test_hierarchy_table_under_named_heading_is_skipped(synthetic_blocks):
    hierarchy = next(
        b for b in synthetic_blocks if b["heading"] == "CIDOC CRM Class Hierarchy"
    )
    assert "CRM Entity" not in hierarchy["text"]


def test_declarations_are_keyed_by_concept_id(synthetic_blocks):
    decl = {b["concept_id"]: b for b in synthetic_blocks if b["kind"] == "declaration"}
    assert set(decl) == {"E1", "E55", "P2"}
    assert decl["E1"]["heading"] == "E1 CRM Entity"
    assert decl["P2"]["heading"] == "P2 has type (is type of)"


def test_declaration_section_path_is_the_enclosing_heading_stack(synthetic_blocks):
    decl = {b["concept_id"]: b for b in synthetic_blocks if b["kind"] == "declaration"}
    assert decl["E1"]["section_path"] == ["CIDOC CRM Class Declarations"]


def test_gap_properties_header_under_normal_style_is_captured(synthetic_blocks):
    """The `Properties:` header styled Normal must not be dropped, and
    everything after it (up to the next boundary) must still be attached to
    the declaration -- not orphaned into a separate/lost block."""
    decl = {b["concept_id"]: b for b in synthetic_blocks if b["kind"] == "declaration"}
    assert "Properties:" in decl["E1"]["text"]
    assert "P2 has type" in decl["E1"]["text"]


def test_gap_dot_one_property_style_is_captured(synthetic_blocks):
    decl = {b["concept_id"]: b for b in synthetic_blocks if b["kind"] == "declaration"}
    assert "P3.1" in decl["E1"]["text"]


def test_declaration_carries_its_fol_line(synthetic_blocks):
    decl = {b["concept_id"]: b for b in synthetic_blocks if b["kind"] == "declaration"}
    assert "E55(x) => E28(x)" in decl["E55"]["text"]


# --- load_document dispatch -------------------------------------------------


def test_load_document_dispatches_docx_by_suffix(tmp_path):
    path = make_synthetic_docx(tmp_path)
    spec = {"id": "t", "path": path.name, "title": "T", "cite": "T"}
    blocks = load_document(spec, root=tmp_path)
    assert any(b["concept_id"] == "E55" for b in blocks)


def test_load_document_honors_the_root_argument(tmp_path):
    """The function must resolve spec['path'] against the *given* root, not
    some hardcoded project path -- otherwise a test pointed at tmp_path would
    silently read the real project file instead."""
    make_synthetic_docx(tmp_path)
    from docx.opc.exceptions import PackageNotFoundError

    other_root = tmp_path / "elsewhere"
    other_root.mkdir()
    spec = {"id": "t", "path": "synthetic.docx", "title": "T", "cite": "T"}
    with pytest.raises(PackageNotFoundError):
        load_document(spec, root=other_root)


def test_load_document_raises_on_unknown_suffix(tmp_path):
    unknown = tmp_path / "doc.pdf"
    unknown.write_bytes(b"%PDF-fake")
    spec = {"id": "t", "path": "doc.pdf", "title": "T", "cite": "T"}
    with pytest.raises(ValueError):
        load_document(spec, root=tmp_path)


# --- chunk_document ----------------------------------------------------------


SPEC = {"id": "test732", "path": "synthetic.docx", "title": "Test Doc", "cite": "Test v1"}


def test_chunk_document_declarations_stay_whole_and_keyed(synthetic_blocks):
    records = chunk_document(SPEC, synthetic_blocks, 2000, 200, r"\b([EP]\d+)\b", FAKE_ONTO)
    e1 = next(r for r in records if r["chunk_id"] == "test732#E1")
    assert e1["kind"] == "declaration"
    assert e1["concept_id"] == "E1"
    assert "P3.1" in e1["text"]  # the whole declaration, unsplit


def test_chunk_document_narrative_gets_sequential_stable_ids(synthetic_blocks):
    records = chunk_document(SPEC, synthetic_blocks, 2000, 200, r"\b([EP]\d+)\b", FAKE_ONTO)
    narrative_ids = [r["chunk_id"] for r in records if r["kind"] == "narrative"]
    assert narrative_ids == sorted(narrative_ids)
    assert all(cid.startswith("test732#s") for cid in narrative_ids)
    assert narrative_ids[0] == "test732#s0000"


def test_chunk_document_splits_narrative_over_chunk_size(synthetic_blocks):
    records = chunk_document(SPEC, synthetic_blocks, 200, 20, r"\b([EP]\d+)\b", FAKE_ONTO)
    big = [r for r in records if r["heading"] == "Big Section"]
    assert len(big) > 1
    tolerance = 200 * 1.1
    assert all(len(r["text"]) <= tolerance for r in big)
    assert all(r["section_path"] == ["Big Section"] for r in big)


def test_chunk_document_narrative_has_null_concept_id(synthetic_blocks):
    records = chunk_document(SPEC, synthetic_blocks, 2000, 200, r"\b([EP]\d+)\b", FAKE_ONTO)
    assert all(r["concept_id"] is None for r in records if r["kind"] == "narrative")


def test_chunk_document_extracts_entities_via_extract_entities(synthetic_blocks):
    records = chunk_document(SPEC, synthetic_blocks, 2000, 200, r"\b([EP]\d+)\b", FAKE_ONTO)
    e1 = next(r for r in records if r["chunk_id"] == "test732#E1")
    # P2 and E55 both appear in E1's text and are "known" per FAKE_ONTO.
    assert "P2" in e1["entities"]
    assert "E55" in e1["entities"]


def test_chunk_document_record_shape_matches_spec(synthetic_blocks):
    records = chunk_document(SPEC, synthetic_blocks, 2000, 200, r"\b([EP]\d+)\b", FAKE_ONTO)
    e55 = next(r for r in records if r["chunk_id"] == "test732#E55")
    assert e55 == {
        "chunk_id": "test732#E55",
        "doc_id": "test732",
        "doc_title": "Test Doc",
        "cite": "Test v1",
        "kind": "declaration",
        "concept_id": "E55",
        "section_path": ["CIDOC CRM Class Declarations"],
        "heading": "E55 Type",
        "text": e55["text"],
        "entities": e55["entities"],
        "entities_historical": e55["entities_historical"],
    }
    assert "E55 comprises controlled-vocabulary terms." in e55["text"]


# --- build_documents ---------------------------------------------------------


def test_build_documents_never_touches_real_project_data(tmp_path):
    """The function must read spec['path'] under the root it is given, not
    under PROJECT_ROOT -- a test passing tmp_path must never fall through to
    the real cidoc_crm_version_7.3.2.docx."""
    make_synthetic_docx(tmp_path)
    cfg = {
        "documents": [
            {"id": "test732", "path": "synthetic.docx", "title": "Test Doc", "cite": "Test v1"}
        ],
        "chunk_size": 2000,
        "chunk_overlap": 200,
        "ontology": {"id_pattern": r"\b([EP]\d+)\b"},
    }
    records = build_documents(cfg, FAKE_ONTO, root=tmp_path)
    assert any(r["chunk_id"] == "test732#E55" for r in records)


def test_build_documents_with_no_documents_configured_returns_empty(tmp_path):
    cfg = {"documents": [], "chunk_size": 2000, "chunk_overlap": 200,
           "ontology": {"id_pattern": r"\b([EP]\d+)\b"}}
    assert build_documents(cfg, FAKE_ONTO, root=tmp_path) == []


# --- integration tests against the real document (read-only) ---------------


@pytest.fixture(scope="module")
def real_blocks():
    if not REAL_DOCX.exists():
        pytest.skip(f"{REAL_DOCX} not present")
    return parse_docx(REAL_DOCX)


@pytest.fixture(scope="module")
def real_onto():
    from lib.ontology import parse_ontology

    cfg = load_config("crm-sig")
    return parse_ontology(PROJECT_ROOT / cfg["ontology"]["xml"])


def test_real_docx_yields_244_declarations_all_with_concept_id(real_blocks):
    decl = [b for b in real_blocks if b["kind"] == "declaration"]
    assert len(decl) == 244
    assert all(b["concept_id"] for b in decl)


def test_real_docx_e55_p2_e1_each_produce_exactly_one_declaration(real_blocks):
    decl = [b for b in real_blocks if b["kind"] == "declaration"]
    for cid in ("E55", "P2", "E1"):
        assert sum(1 for b in decl if b["concept_id"] == cid) == 1, cid


def test_real_docx_e55_declaration_has_scope_note_and_fol_line(real_blocks):
    decl = {b["concept_id"]: b for b in real_blocks if b["kind"] == "declaration"}
    text = decl["E55"]["text"]
    assert "concepts denoted by terms from thesauri" in text  # scope note
    assert "E55(x)" in text  # FOL line


def test_real_docx_e1_declaration_carries_its_dot_one_properties(real_blocks):
    """E1 owns P3 and P137, whose `.1` properties-of-properties (style
    `CRM Dot One Property`) are declared right beneath them -- this is the
    concrete instance of gap 2 the design note calls out."""
    decl = {b["concept_id"]: b for b in real_blocks if b["kind"] == "declaration"}
    text = decl["E1"]["text"]
    assert "P3.1" in text
    assert "P137.1" in text


def test_real_docx_narrative_chunks_all_carry_nonempty_section_path(real_blocks, real_onto):
    cfg = load_config("crm-sig")
    spec = {"id": "crm732", "path": REAL_DOCX.name, "title": "t", "cite": "c"}
    records = chunk_document(
        spec, real_blocks, cfg["chunk_size"], cfg["chunk_overlap"],
        cfg["ontology"]["id_pattern"], real_onto,
    )
    narrative = [r for r in records if r["kind"] == "narrative"]
    assert narrative
    assert all(r["section_path"] for r in narrative)


def test_real_docx_narrative_chunks_respect_chunk_size_tolerance(real_blocks, real_onto):
    cfg = load_config("crm-sig")
    spec = {"id": "crm732", "path": REAL_DOCX.name, "title": "t", "cite": "c"}
    records = chunk_document(
        spec, real_blocks, cfg["chunk_size"], cfg["chunk_overlap"],
        cfg["ontology"]["id_pattern"], real_onto,
    )
    narrative = [r for r in records if r["kind"] == "narrative"]
    tolerance = cfg["chunk_size"] * 1.1
    over = [len(r["text"]) for r in narrative if len(r["text"]) > tolerance]
    assert over == []


def test_real_docx_modelling_principles_minimality_survives(real_blocks, real_onto):
    cfg = load_config("crm-sig")
    spec = {"id": "crm732", "path": REAL_DOCX.name, "title": "t", "cite": "c"}
    records = chunk_document(
        spec, real_blocks, cfg["chunk_size"], cfg["chunk_overlap"],
        cfg["ontology"]["id_pattern"], real_onto,
    )
    hits = [
        r for r in records
        if r["section_path"][:2] == ["Modelling principles", "Minimality"]
        and "economically as possible" in r["text"]
    ]
    assert hits


def test_real_docx_total_narrative_chars_within_5_pct_of_measured_baseline(real_blocks):
    """Catches silent drops in the parser -- measured against the prototype's
    191,829 chars across 89 narrative blocks. Checked on the raw parsed
    blocks (pre-chunking), since splitting with overlap would inflate the
    total and mask a real drop."""
    narrative = [b for b in real_blocks if b["kind"] == "narrative"]
    total = sum(len(b["text"]) for b in narrative)
    baseline = 191_829
    assert abs(total - baseline) / baseline < 0.05, total


def test_a_document_may_declare_its_own_kind():
    """The default is "narrative", which lib.retrieve searches by default as
    part of the reference model. Correct for the specification, wrong for a
    working draft sitting beside it: indexed as narrative, the Conceptual
    Modelling Principles' 66 short chunks -- each carrying the word
    "principle" in an "ID | Principle | Slogan" header row -- took all five
    slots for "minimality principle" and pushed the specification's own
    Minimality section off the page. Same burial that put SPEC_KINDS in
    lib/retrieve.py, one level down."""
    from lib.documents import chunk_document

    blocks = [{"kind": "narrative", "concept_id": None, "section_path": ["A"],
               "heading": "A", "text": "some guidance text"}]
    spec = {"id": "d", "title": "t", "cite": "c", "kind": "principles"}
    out = chunk_document(spec, blocks, 2000, 200, r"\b([EP]\d{1,3})\b",
                         {"classes": {}, "properties": {}, "historical": {}})
    assert [r["kind"] for r in out] == ["principles"]

    default = chunk_document({"id": "d", "title": "t", "cite": "c"}, blocks,
                             2000, 200, r"\b([EP]\d{1,3})\b",
                             {"classes": {}, "properties": {}, "historical": {}})
    assert [r["kind"] for r in default] == ["narrative"]


def test_reference_apparatus_is_excluded_from_document_search():
    # Bibliographies are 43 of 7,086 chunks (0.6%) and took two of the top
    # five slots on "grave goods burial ship burial archaeology", because a
    # Works Cited entry is a dense list of proper nouns and that is what BM25
    # rewards. An agent modelling the Sutton Hoo helmet hit exactly this and
    # reported the tool as useless.
    from lib.retrieve import _is_apparatus

    assert _is_apparatus({"section_path": ["Works Cited"]})
    assert _is_apparatus({"section_path": ["Introduction", "Further Reading"]})
    assert _is_apparatus({"heading": "See also"})
    # Real content must not be caught by the same net.
    assert not _is_apparatus({"section_path": ["Introduction", "Applied Form",
                                               "Property Quantifiers"]})
    assert not _is_apparatus({"heading": "E22 Human-Made Object"})
