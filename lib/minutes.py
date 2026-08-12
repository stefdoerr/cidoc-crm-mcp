"""Parse the CIDOC CRM SIG meeting minutes into searchable chunks.

The minutes are the fourth corpus here, and the only one that records
*deliberation*. The issue register states that issue 295 deleted E84 at
Cologne in January 2018; the mailing list argues towards it; the minutes for
that meeting are where the room decided. A question like "why was E84
removed" is answerable from the register alone only as a fact, and from the
minutes as a reason.

Three input formats, all of which had to be supported rather than picked
between, because the archive is not evenly published:

  * **docx** (31 meetings) -- authored format, and the only one carrying
    paragraph styles. Modern minutes put each agenda item under a `Heading 2`
    reading "ISSUE 345: properties having domain or range deprecated classes",
    which is exactly the boundary a chunk wants.
  * **pdf** (24 meetings) -- no styles, so headings are recovered from the
    text. Page furniture (numbers, running heads) comes through as inline
    text; `_strip_page_furniture` removes the lines that repeat across pages.
  * **doc** (15 meetings) -- Word 97. Needs the FIB piece table, implemented
    below, because no converter is available in this environment and these
    are not optional: they are meetings 17-30 and the early Chios and FRBR
    harmonisation meetings, where E27, E84 and P39 were settled.

Chunking follows the issue headings where they exist, because an agenda item
is the natural unit and it maps one-to-one onto the register. Measured on
this corpus: 62 of 70 meetings expose at least one issue heading, and the
eight that do not are genuinely pre-register -- Stavanger 1995, Nuremberg
1997, London 1999, Ottawa 2000 and the first harmonisation meeting, all
predating numbered issues -- plus one position paper filed among the
minutes. Those fall back to size-bounded blocks.

Recognising the headings took two passes, and the first one silently lost a
third of the corpus's decisions. Only the modern minutes write "ISSUE 345:";
the 2010-2013 meetings head each item with a bare number -- "204: The issue
is done" -- and matching only the explicit form left 26 meetings looking
issue-free. Adding the bare form (validated against the register, never on
shape alone) took issue-bearing chunks from 1,045 to 1,891 and the meetings
with no issue heading from 26 to 8.

`MAX_CHUNK_CHARS` is not decoration. Issue-page ingestion previously produced
15 chunks over 20,000 characters and ran the embedding model out of memory;
these minutes reach 100,000 characters for a single meeting and one agenda
item can run to several thousand, so every section is split to fit whether or
not its heading structure suggests it should be.
"""

import hashlib
import re
import struct
from pathlib import Path

from lib.config import DATA_DIR

MINUTES_DIR = DATA_DIR / "minutes"

# Held well under the embedding model's limit; see the module docstring for
# what happens when a corpus is trusted to be self-limiting.
MAX_CHUNK_CHARS = 6_000
MIN_CHUNK_CHARS = 120

# "ISSUE 345:", "Issue 161:How to organize extensions", "Issue 172, 173",
# "Issue 166" alone on a line. The number must be adjacent to the word, the
# same rule lib.issues uses -- see that module for why a looser gap pulls in
# dates and CRM identifiers.
ISSUE_HEADING = re.compile(r"^\s*issues?\s*[:#]?\s*(\d{1,4})\b", re.I)
# The run of numbers immediately after the keyword, so "Issue 172, 173" links
# to both. Reading instead "everything before the first colon" looks
# equivalent and is not: "issue: 204" puts the colon before the number, so
# that rule inspects the word "issue" alone and finds nothing.
_HEADING_NUMBER_RUN = re.compile(
    r"^\s*issues?\s*[:#]?\s*((?:\d{1,4})(?:\s*(?:,|&|and|/)\s*\d{1,4})*)", re.I)
_HEADING_NUMBERS = re.compile(r"\d{1,4}")

# The 2010-2013 minutes drop the word entirely and head each item with a bare
# number: "204: The issue is done", "207: The proposal is accepted, the issue
# was closed." Missing these was not a cosmetic loss -- that era records its
# outcomes more tersely and more decisively than any other, and eleven
# meetings looked issue-free until this pattern was added.
#
# A bare "204:" is far too weak a signal on its own; it matches numbered
# lists, clock times and section numbers. So it counts only when the number
# is one the SIG actually assigned, checked against the real register. That
# is the same collection-as-authority rule lib.issues.validate_issue_ids
# applies to prose mentions, for the same reason: shape is not authority.
_BARE_NUMBER_HEADING = re.compile(r"^\s*(\d{1,4})\s*[:.)]\s*\S")

_MEETING_TITLE = re.compile(
    r"\b(\d{1,3})(?:st|nd|rd|th)?\s+(?:joint\s+)?"
    r"(?:meeting|CIDOC|SIG)\b", re.I)
_MONTH = ("January|February|March|April|May|June|July|August|September|"
          "October|November|December")
# Tried in order, most specific first, against the header paragraphs JOINED
# into one string rather than line by line. Joining is what matters: a PDF
# breaks "9-10th December\n2004" across two lines, and a per-line search sees
# a month with no year and a year with no month. 21 of 70 meetings dated only
# once these three forms and the join were in place.
_DATE_PATTERNS = (
    # 27 - 30 November, 2018  |  9-10th December 2004  |  1 May 2009
    re.compile(rf"\b(?:\d{{1,2}}(?:st|nd|rd|th)?\s*[-–]\s*)?"
               rf"\d{{1,2}}(?:st|nd|rd|th)?\s+(?:{_MONTH})\s*,?\s*(?:19|20)\d{{2}}\b", re.I),
    # December 2004  |  November, 2018
    re.compile(rf"\b(?:{_MONTH})\s*,?\s*(?:19|20)\d{{2}}\b", re.I),
    # 15/9/2008  |  15.09.2008  |  2008-09-15
    re.compile(r"\b\d{1,2}[/.]\d{1,2}[/.](?:19|20)\d{2}\b"),
    re.compile(r"\b(?:19|20)\d{2}-\d{1,2}-\d{1,2}\b"),
)
# Last resort. Restricted to the header window, because a bare year matched
# against the body picks up the first citation or scope-note example instead
# of the meeting -- which is worse than reporting no date at all.
_BARE_YEAR = re.compile(r"\b(19[89]\d|20[0-4]\d)\b")


# --------------------------------------------------------------------------
# text extraction
# --------------------------------------------------------------------------

def _docx_paragraphs(path: Path) -> list[tuple[str | None, str]]:
    """(style name, text) for every paragraph and table cell.

    Tables carry real content in these minutes -- proposed scope notes are
    routinely laid out as before/after tables -- so skipping them would drop
    the substance of the decision and keep only the preamble.
    """
    import docx

    document = docx.Document(str(path))
    out: list[tuple[str | None, str]] = []
    for para in document.paragraphs:
        text = " ".join(para.text.split())
        if text:
            out.append((para.style.name if para.style is not None else None, text))
    for table in document.tables:
        for row in table.rows:
            cells = [" ".join(c.text.split()) for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                out.append(("Table", line))
    return out


def _doc_text(path: Path) -> str:
    """Word 97-2003 body text, via the FIB piece table.

    The piece table is the documented way and the only reliable one: a
    fastsaved .doc stores its text in fragments that are not in reading order
    on disk, so scanning the WordDocument stream for printable runs returns
    the paragraphs shuffled. Each piece declares whether it is CP1252
    (compressed, bit 30 of fc set, offset halved) or UTF-16LE.

    Validated against ground truth: `Meeting10_Minutes.doc` and the same
    meeting's PDF, where available, agree on the body text.
    """
    import olefile

    with olefile.OleFileIO(str(path)) as ole:
        word = ole.openstream("WordDocument").read()
        # fWhichTblStm (bit 9 of the FIB flags) says which table stream is live.
        which = (struct.unpack_from("<H", word, 0x000A)[0] >> 9) & 1
        names = {"/".join(entry) for entry in ole.listdir()}
        table_name = "1Table" if which else "0Table"
        if table_name not in names:
            table_name = "0Table" if which else "1Table"
        if table_name not in names:
            raise ValueError(f"{path.name}: no table stream")
        table = ole.openstream(table_name).read()

    fc_clx, lcb_clx = struct.unpack_from("<II", word, 0x01A2)
    clx = table[fc_clx:fc_clx + lcb_clx]
    # Skip any Prc entries (type 0x01) to reach the Pcdt (type 0x02).
    i = 0
    while i < len(clx) and clx[i] == 1:
        i += 3 + struct.unpack_from("<h", clx, i + 1)[0]
    if i >= len(clx) or clx[i] != 2:
        raise ValueError(f"{path.name}: no piece table in CLX")
    lcb_pcdt = struct.unpack_from("<I", clx, i + 1)[0]
    pcdt = clx[i + 5:i + 5 + lcb_pcdt]

    count = (len(pcdt) - 4) // 12
    cps = [struct.unpack_from("<I", pcdt, 4 * k)[0] for k in range(count + 1)]
    parts = []
    for k in range(count):
        base = 4 * (count + 1) + 8 * k
        fc = struct.unpack_from("<I", pcdt, base + 2)[0]
        compressed = bool(fc & 0x40000000)
        offset = (fc & 0x3FFFFFFF) // 2 if compressed else (fc & 0x3FFFFFFF)
        length = cps[k + 1] - cps[k]
        if compressed:
            parts.append(word[offset:offset + length].decode("cp1252", "replace"))
        else:
            parts.append(word[offset:offset + length * 2].decode("utf-16-le", "replace"))
    return "".join(parts).replace("\r", "\n").replace("\x07", "\n")


def _strip_page_furniture(pages: list[str]) -> list[str]:
    """Drop lines that repeat on most pages -- running heads and page numbers.

    A PDF of minutes reproduces its header on every page, and left in place
    that text is indexed dozens of times per meeting and outranks the content
    on any query matching the meeting's own title. Only lines appearing on
    more than half the pages are removed, so a genuinely repeated sentence in
    a short document survives.
    """
    if len(pages) < 3:
        return pages
    counts: dict[str, int] = {}
    for page in pages:
        for line in {" ".join(x.split()) for x in page.split("\n") if x.strip()}:
            if len(line) < 90:
                counts[line] = counts.get(line, 0) + 1
    boilerplate = {line for line, n in counts.items() if n > len(pages) * 0.5}
    cleaned = []
    for page in pages:
        keep = [x for x in page.split("\n")
                if " ".join(x.split()) not in boilerplate
                and not re.fullmatch(r"\s*\d{1,3}\s*", x)]
        cleaned.append("\n".join(keep))
    return cleaned


def _pdf_paragraphs(path: Path) -> list[tuple[str | None, str]]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [(page.extract_text() or "") for page in reader.pages]
    out = []
    for page in _strip_page_furniture(pages):
        for line in page.split("\n"):
            text = " ".join(line.split())
            if text:
                out.append((None, text))
    return out


def extract_paragraphs(path: str | Path) -> list[tuple[str | None, str]]:
    """(style, text) pairs for any supported minutes file.

    `style` is a real Word style name for docx, "Table" for table rows, and
    None for the formats that carry no styling. Callers must treat None as
    "unknown", never as "body text": a PDF heading is a heading, it just
    cannot say so.
    """
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _docx_paragraphs(path)
    if suffix == ".pdf":
        return _pdf_paragraphs(path)
    if suffix == ".doc":
        return [(None, " ".join(line.split()))
                for line in _doc_text(path).split("\n") if line.strip()]
    raise ValueError(f"unsupported minutes format: {path.name}")


# --------------------------------------------------------------------------
# structure
# --------------------------------------------------------------------------

def is_heading(style: str | None, text: str,
               known_issues: set[int] | None = None) -> bool:
    """Whether this paragraph opens a new section.

    A Word heading style settles it. Otherwise the text has to look like one,
    and only two patterns are trusted: an explicit issue heading, and a bare
    leading number that the register confirms is a real issue (see
    `_BARE_NUMBER_HEADING`). Guessing from length or capitalisation misfires
    badly on the participant lists and quoted scope notes these documents are
    full of.

    Without `known_issues` the bare-number form is not recognised at all,
    which is the safe default: better to under-segment than to cut a document
    on every numbered list item.
    """
    if style and style.startswith("Heading"):
        return True
    if ISSUE_HEADING.match(text):
        return True
    return bool(heading_issues(text, known_issues))


def heading_issues(text: str, known_issues: set[int] | None = None) -> list[int]:
    """Issue numbers named in a heading: "Issue 172, 173" -> [172, 173].

    For the explicit form, every number before the colon is an issue number,
    which is why this may be looser than lib.issues' sieve without letting
    dates or CRM ids through. For the bare-number form the register is the
    authority and an unknown number yields nothing.
    """
    run = _HEADING_NUMBER_RUN.match(text)
    if run:
        return [int(n) for n in _HEADING_NUMBERS.findall(run.group(1))][:8]
    if known_issues:
        match = _BARE_NUMBER_HEADING.match(text)
        if match and int(match.group(1)) in known_issues:
            return [int(match.group(1))]
    return []


def parse_header(paragraphs: list[tuple[str | None, str]]) -> dict:
    """Meeting title, date and venue from the first few paragraphs.

    Best-effort by construction: 70 meetings across 30 years have no common
    template, and a missing field is reported as None rather than guessed.
    The title is whichever of the opening lines names a meeting number, which
    is more reliable than "the first line" -- several files open with a logo
    caption or a document-control line.
    """
    head = [text for _, text in paragraphs[:25]]
    title = next((t for t in head if _MEETING_TITLE.search(t) and len(t) > 15), None)
    if title is None:
        title = next((t for t in head if len(t) > 15), None)

    joined = " ".join(head)
    date = None
    for pattern in _DATE_PATTERNS:
        match = pattern.search(joined)
        if match:
            date = " ".join(match.group(0).split())
            break
    if date is None:
        match = _BARE_YEAR.search(joined)
        date = match.group(0) if match else None

    venue = None
    if title is not None:
        after = head[head.index(title) + 1:head.index(title) + 5]
        venue = next((t for t in after
                      if not any(p.search(t) for p in _DATE_PATTERNS)
                      and 3 < len(t) < 120), None)
    return {"title": title, "date": date, "venue": venue}


def _split_to_size(text: str, limit: int = MAX_CHUNK_CHARS) -> list[str]:
    """Split on sentence-ish boundaries so no chunk exceeds `limit`."""
    if len(text) <= limit:
        return [text]
    out, current = [], ""
    for piece in re.split(r"(?<=[.!?])\s+", text):
        while len(piece) > limit:            # a single unpunctuated wall of text
            out.append(piece[:limit])
            piece = piece[limit:]
        if len(current) + len(piece) + 1 > limit:
            if current:
                out.append(current.strip())
            current = piece
        else:
            current = f"{current} {piece}".strip()
    if current.strip():
        out.append(current.strip())
    return out


def sections(paragraphs: list[tuple[str | None, str]],
             known_issues: set[int] | None = None) -> list[dict]:
    """Group paragraphs under their headings.

    Content before the first heading becomes a "Preamble" section rather than
    being dropped: that is where the participant list and the meeting's own
    date and venue live, and questions about who attended are legitimate.

    A heading with no body still yields a section. Minutes entries are often
    a single clause -- "332: Issue closed", "It is closed." -- and those are
    the most decision-bearing lines in the corpus, not noise to be dropped
    for being short.
    """
    out: list[dict] = []
    current = {"heading": "Preamble", "issues": [], "lines": []}
    for style, text in paragraphs:
        if is_heading(style, text, known_issues):
            if current["lines"] or current["issues"]:
                out.append(current)
            current = {"heading": text,
                       "issues": heading_issues(text, known_issues),
                       "lines": []}
        else:
            current["lines"].append(text)
    if current["lines"] or current["issues"]:
        out.append(current)
    return out


def load_known_issues(path: str | Path | None = None) -> set[int]:
    """Every issue number the SIG ever assigned, from the committed register.

    The authority for the bare-number heading form. Returns an empty set if
    the register is absent, which disables that form rather than guessing.
    """
    import json

    path = Path(path) if path else DATA_DIR.parent / "crm_issues.json"
    if not path.exists():
        return set()
    return {int(k) for k in json.loads(path.read_text(encoding="utf-8"))["entries"]}


def build_minutes_chunks(path: str | Path, doc_id: str | None = None,
                         known_issues: set[int] | None = None) -> list[dict]:
    """One meeting -> chunk records in data/documents.jsonl's schema.

    `kind` is "minutes", alongside "declaration", "narrative" and "issue".
    `issues` carries the issue numbers this chunk is about, taken from its
    heading, which is what lets `search.py issue <n>` show the meetings that
    discussed it without re-scanning the prose.
    """
    path = Path(path)
    doc_id = doc_id or path.stem
    paragraphs = extract_paragraphs(path)
    if not paragraphs:
        return []
    header = parse_header(paragraphs)
    title = header["title"] or doc_id
    cite_date = f", {header['date']}" if header["date"] else ""

    chunks = []
    for index, section in enumerate(sections(paragraphs, known_issues)):
        body = "\n".join(section["lines"]).strip()
        # A heading with no body is a real record -- the meeting reached that
        # agenda item -- but it is not retrievable content, and indexing
        # "ISSUE 380" as a chunk puts a bare identifier into the vector store
        # to compete with the discussion of issue 380 elsewhere. The linkage
        # is preserved by `issue_links`, which reads sections directly.
        if not body:
            continue
        if len(body) < MIN_CHUNK_CHARS and not section["issues"]:
            continue
        # The heading leads the indexed text. Many issue entries are a single
        # clause ("Issue closed", "CEO presented his HW") and a chunk holding
        # only those words is unretrievable and unreadable out of context --
        # it does not say which issue it closed. Keeping the heading in the
        # body makes the chunk self-describing to both the reader and the
        # embedding, at the cost of repeating it in the metadata.
        head = "" if section["heading"] == "Preamble" else section["heading"]
        budget = max(MIN_CHUNK_CHARS, MAX_CHUNK_CHARS - len(head) - 1)
        for part, piece in enumerate(_split_to_size(body, budget)):
            piece = f"{head}\n{piece}".strip() if head else piece
            suffix = f"-{part}" if part else ""
            chunks.append({
                "chunk_id": f"minutes:{doc_id}#s{index:03d}{suffix}",
                "doc_id": doc_id,
                "doc_title": title,
                "kind": "minutes",
                "heading": section["heading"],
                "section_path": [title, section["heading"]],
                "cite": f"{title}{cite_date} — {section['heading']}",
                "meeting_date": header["date"],
                "meeting_venue": header["venue"],
                "issues": section["issues"],
                "text": piece,
            })
    return chunks


def issue_links(path: str | Path, doc_id: str | None = None,
                known_issues: set[int] | None = None) -> list[dict]:
    """Which issues this meeting took up, one row per (issue, agenda item).

    Deliberately separate from `build_minutes_chunks`, and computed from the
    sections rather than the chunks, so an agenda item that the minutes
    record without discussion ("ISSUE 380" and nothing else, or "332: Issue
    closed") still links the meeting to the issue. Those are exactly the
    entries a question like "when was issue 332 closed" needs, and exactly
    the ones too short to survive as retrievable chunks.
    """
    path = Path(path)
    doc_id = doc_id or path.stem
    paragraphs = extract_paragraphs(path)
    if not paragraphs:
        return []
    header = parse_header(paragraphs)
    rows = []
    for index, section in enumerate(sections(paragraphs, known_issues)):
        body = "\n".join(section["lines"]).strip()
        for issue in section["issues"]:
            rows.append({
                "issue": issue,
                "doc_id": doc_id,
                "title": header["title"] or doc_id,
                "date": header["date"],
                "heading": section["heading"],
                "chunk_id": f"minutes:{doc_id}#s{index:03d}" if body else None,
                "chars": len(body),
            })
    return rows


SUPPORTED_SUFFIXES = (".docx", ".pdf", ".doc")


def load_all(minutes_dir: str | Path | None = None,
             known_issues: set[int] | None = None) -> dict:
    """Every cached meeting -> chunks, issue links, and what could not be read.

    Unreadable files are reported, never skipped silently. One real file in
    this corpus proves the need: the 61st joint meeting's .docx is malformed
    (python-docx raises `KeyError: 'word/NULL'`), and since the fetcher takes
    only the preferred format, a silent skip would have dropped a 2025 meeting
    with 46 issues on it while every count still looked plausible. The
    `unreadable` list is what makes that visible; the fix is to fetch the
    alternate rendering, which the index records.

    Content-duplicate meetings are detected here rather than assumed absent --
    see `content_key` for why filenames cannot settle it.
    """
    minutes_dir = Path(minutes_dir) if minutes_dir else MINUTES_DIR
    if known_issues is None:
        known_issues = load_known_issues()

    chunks: list[dict] = []
    links: list[dict] = []
    unreadable: list[dict] = []
    keys: dict[str, str] = {}

    files = [p for p in sorted(minutes_dir.iterdir())
             if p.suffix.lower() in SUPPORTED_SUFFIXES]
    for path in files:
        try:
            paragraphs = extract_paragraphs(path)
        except Exception as exc:                       # noqa: BLE001 - reported
            unreadable.append({"file": path.name, "error": f"{type(exc).__name__}: {exc}"})
            continue
        if not paragraphs:
            unreadable.append({"file": path.name, "error": "no text extracted"})
            continue
        keys[path.stem] = content_key(paragraphs)
        chunks.extend(build_minutes_chunks(path, known_issues=known_issues))
        links.extend(issue_links(path, known_issues=known_issues))

    return {
        "chunks": chunks,
        "issue_links": links,
        "meetings": len(files) - len(unreadable),
        "files": len(files),
        "unreadable": unreadable,
        "duplicates": duplicate_groups(keys),
    }


# --------------------------------------------------------------------------
# duplicate meetings
# --------------------------------------------------------------------------

def content_key(paragraphs: list[tuple[str | None, str]], sample: int = 4000) -> str:
    """A hash of the document's own words, for spotting the same meeting
    published under two filenames.

    Deliberately not a hash of the file: the same minutes as .docx and as .pdf
    are different bytes and different lengths. Deliberately not the parsed
    meeting number either -- `Meeting10_Minutes.doc` is the 10th FRBR/CRM
    harmonisation meeting and `10th_crm_meeting_minutes.pdf` is the 10th CRM
    SIG, four years and one numbering series apart, and merging them on the
    number would silently delete a meeting. Only the text can tell.
    """
    words = " ".join(text for _, text in paragraphs).lower()
    words = re.sub(r"[^a-z0-9 ]+", " ", words)
    words = " ".join(words.split())[:sample]
    return hashlib.sha256(words.encode("utf-8")).hexdigest()


def duplicate_groups(keyed: dict[str, str]) -> list[list[str]]:
    """{doc_id: content_key} -> groups of doc_ids sharing a key, size > 1."""
    buckets: dict[str, list[str]] = {}
    for doc_id, key in keyed.items():
        buckets.setdefault(key, []).append(doc_id)
    return [sorted(v) for v in buckets.values() if len(v) > 1]
